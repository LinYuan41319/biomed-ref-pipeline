import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document

from biomed_ref_pipeline.zotero_ooxml import insert_zotero_fields


class ZoteroOoxmlTests(unittest.TestCase):
    def test_insert_zotero_fields(self):
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "source.docx"
            out = Path(tmp) / "zotero.docx"
            doc = Document()
            doc.add_paragraph("Claim (PMID: 12345678, 23456789).")
            doc.save(source)

            groups = [
                {
                    "raw_text": "(PMID: 12345678, 23456789)",
                    "pmids": "12345678 23456789",
                }
            ]
            metadata = [
                {"pmid": "12345678", "authors": "Smith AB", "title": "A", "journal": "J", "year": "2025"},
                {"pmid": "23456789", "authors": "Wang CD", "title": "B", "journal": "K", "year": "2024"},
            ]
            report = insert_zotero_fields(source, groups, metadata, out)
            self.assertEqual(report["citation_field_count"], 1)
            self.assertEqual(report["bibliography_field_count"], 1)
            self.assertEqual(report["remaining_placeholder_count"], 0)

            with zipfile.ZipFile(out) as z:
                xml = z.read("word/document.xml").decode("utf-8")
            self.assertIn("ADDIN ZOTERO_ITEM CSL_CITATION", xml)
            self.assertIn("ADDIN ZOTERO_BIBL", xml)
            self.assertNotIn("(PMID: 12345678, 23456789)", xml)

    def test_insert_zotero_fields_skip_missing(self):
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "source.docx"
            out = Path(tmp) / "zotero.docx"
            doc = Document()
            doc.add_paragraph("Claim (PMID: 12345678).")
            doc.add_paragraph("Missing (PMID: 99999999).")
            doc.save(source)

            groups = [
                {"raw_text": "(PMID: 12345678)", "pmids": "12345678"},
                {"raw_text": "(PMID: 99999999)", "pmids": "99999999"},
            ]
            metadata = [{"pmid": "12345678", "authors": "Smith AB", "title": "A", "journal": "J", "year": "2025"}]
            report = insert_zotero_fields(source, groups, metadata, out, on_missing="skip")
            self.assertEqual(report["citation_field_count"], 1)
            self.assertEqual(report["skipped_group_count"], 1)
            self.assertIn("99999999", report["skipped_groups"][0]["missing_pmids"])


if __name__ == "__main__":
    unittest.main()
