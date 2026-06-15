from __future__ import annotations

import re
from pathlib import Path

from .pmid import extract_pmids_from_text


OVERCLAIM_PATTERNS = [
    (re.compile(r"\b(causally drives|proves|confirms|definitively demonstrates)\b", re.I), "Strong causal wording"),
    (re.compile(r"\b(validated binding|directly binds|confirmed target)\b", re.I), "Binding or target validation wording"),
    (re.compile(r"\b(clinical tool|diagnostic tool|therapeutic target)\b", re.I), "Clinical translation wording"),
]

WEAK_EVIDENCE_CONTEXT = re.compile(
    r"\b(bioinformatic|transcriptomic|proteomic|single-cell|spatial|association|correlation|molecular docking)\b",
    re.I,
)

AI_STYLE_PATTERNS = [
    re.compile(r"\bthis review (will|aims to|seeks to)\b", re.I),
    re.compile(r"\bit is important to (note|consider)\b", re.I),
    re.compile(r"\bfuture studies are needed\b", re.I),
]


def read_docx_text(path: str | Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required to read .docx files") from exc
    doc = Document(str(path))
    parts = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(parts)


def manuscript_qc_report(text: str) -> str:
    findings: list[str] = []
    lines = [line.strip() for line in text.splitlines() if line.strip()]

    for number, line in enumerate(lines, start=1):
        pmid_count = len(extract_pmids_from_text(line))
        if pmid_count > 3:
            findings.append(f"- Line {number}: citation density has {pmid_count} PMID markers; keep one evidence claim to <=3 citations.")

        has_weak_context = bool(WEAK_EVIDENCE_CONTEXT.search(line))
        for pattern, label in OVERCLAIM_PATTERNS:
            if pattern.search(line):
                detail = " This is higher risk because the sentence also contains weak-evidence context." if has_weak_context else ""
                findings.append(f"- Line {number}: {label}.{detail} Text: {line[:240]}")

        for pattern in AI_STYLE_PATTERNS:
            if pattern.search(line):
                findings.append(f"- Line {number}: generic or task-like review phrasing. Text: {line[:240]}")

    if not findings:
        findings.append("- No heuristic QC findings detected.")

    return "\n".join(
        [
            "# Manuscript QC Report",
            "",
            "This report is heuristic. It flags wording that should be reviewed by a domain expert.",
            "",
            "## Findings",
            "",
            *findings,
            "",
        ]
    )
