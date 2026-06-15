from __future__ import annotations

import re
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Iterable


PMID_CONTEXT_RE = re.compile(
    r"(?i)\b(?:pmids?|pubmed\s*ids?)\s*[:#\uFF1A]?\s*([0-9][0-9,\s;\uFF1B\u3001/\uFF0C&+\-andAND]{4,240})"
)
PMID_PAREN_RE = re.compile(r"(?i)\((?:pmids?|pubmed\s*ids?)\s*[:#\uFF1A]?\s*([^)]+)\)")
PMID_DIGIT_RE = re.compile(r"(?<!\d)(\d{6,9})(?!\d)")


@dataclass(frozen=True)
class PmidGroup:
    source: str
    location: str
    raw_text: str
    pmids: list[str]

    def to_row(self) -> dict[str, str]:
        row = asdict(self)
        row["pmids"] = " ".join(self.pmids)
        return row


def extract_pmids_from_text(text: str) -> list[str]:
    pmids: list[str] = []
    seen: set[str] = set()
    for match in PMID_DIGIT_RE.finditer(text):
        value = match.group(1)
        if value not in seen:
            seen.add(value)
            pmids.append(value)
    return pmids


def extract_pmid_groups_from_text(text: str, source: str = "text", location: str = "body") -> list[PmidGroup]:
    groups: list[PmidGroup] = []
    spans: list[tuple[int, int]] = []
    for pattern in (PMID_PAREN_RE, PMID_CONTEXT_RE):
        for match in pattern.finditer(text):
            if any(_overlaps(match.span(), span) for span in spans):
                continue
            raw = match.group(0)
            pmids = extract_pmids_from_text(raw)
            if pmids:
                spans.append(match.span())
                groups.append(PmidGroup(source=source, location=location, raw_text=raw, pmids=pmids))
    return groups


def extract_pmid_groups_from_docx(path: str | Path) -> list[PmidGroup]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required to read .docx files") from exc

    doc = Document(str(path))
    source = str(path)
    groups: list[PmidGroup] = []

    for idx, paragraph in enumerate(doc.paragraphs, start=1):
        groups.extend(extract_pmid_groups_from_text(paragraph.text, source=source, location=f"paragraph:{idx}"))

    for table_index, table in enumerate(doc.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            for cell_index, cell in enumerate(row.cells, start=1):
                cell_text = "\n".join(paragraph.text for paragraph in cell.paragraphs)
                groups.extend(
                    extract_pmid_groups_from_text(
                        cell_text,
                        source=source,
                        location=f"table:{table_index}/row:{row_index}/cell:{cell_index}",
                    )
                )
    return groups


def unique_pmids(groups: Iterable[PmidGroup]) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    for group in groups:
        for pmid in group.pmids:
            if pmid not in seen:
                seen.add(pmid)
                result.append(pmid)
    return result


def _overlaps(a: tuple[int, int], b: tuple[int, int]) -> bool:
    return a[0] < b[1] and b[0] < a[1]
